"""
SIGDOC-ML — Motor de análisis de curriculums v2.0
==================================================
Ponderación:
  30% → Similitud semántica con el perfil (TF-IDF coseno)
  25% → Habilidades técnicas (keywords con frecuencia)
  20% → Años de experiencia laboral
  15% → Nivel educativo alcanzado
  10% → Completitud del CV (secciones presentes)

Salida JSON (sin texto extraído — se guarda en BD por separado):
{
  "puntaje": 87.4,
  "categoria": "APROBADO",
  "nivel": "ALTO",
  "observaciones": "...",
  "texto_extraido": "...",   ← solo para la BD, no para el ranking
  "detalles": {
    "perfil_pct":       85.0,
    "habilidades_pct":  100.0,
    "experiencia_pct":  90.0,
    "educacion_pct":    80.0,
    "completitud_pct":  95.0,
    "anios_experiencia": 6,
    "nivel_educativo": "licenciatura",
    "habilidades_encontradas": ["php", "mysql"],
    "habilidades_total": 3,
    "idiomas": ["ingles"],
    "secciones": {"experiencia": true, "educacion": true, ...}
  }
}

Modos de uso:
  python analizar_cv.py <ruta_cv> <perfil> <keywords>
  python analizar_cv.py --args-file <ruta_json>

Dependencias:
  pip install scikit-learn pdfminer.six python-docx unidecode
"""

import sys
import json
import os
import re
import unicodedata
import logging

# Silenciar warnings de pdfminer en stderr
logging.getLogger("pdfminer").setLevel(logging.ERROR)

# ══════════════════════════════════════════════════════════════
# TABLAS DE PESOS
# ══════════════════════════════════════════════════════════════

NIVELES_EDUCATIVOS = {
    "doctorado":    100,
    "phd":          100,
    "maestria":      90,
    "master":        90,
    "mba":           90,
    "especialidad":  85,
    "licenciatura":  80,
    "licenciado":    80,
    "bachiller":     70,
    "ingeniero":     75,
    "abogado":       75,
    "medico":        75,
    "contador":      70,
    "tecnico":       50,
    "tecnólogo":     50,
    "tecnologo":     50,
    "secundaria":    20,
    "bachillerato":  20,
}

IDIOMAS_CONOCIDOS = {
    "ingles":   "Inglés",
    "english":  "Inglés",
    "frances":  "Francés",
    "french":   "Francés",
    "aleman":   "Alemán",
    "german":   "Alemán",
    "portugues":"Portugués",
    "portuguese":"Portugués",
    "italiano": "Italiano",
    "chinese":  "Chino",
    "chino":    "Chino",
    "japones":  "Japonés",
    "japanese": "Japonés",
}

CATEGORIAS = [
    (90, "EXCELENTE",      "MUY ALTO"),
    (75, "APROBADO",       "ALTO"),
    (60, "A CONSIDERAR",   "MEDIO"),
    (0,  "NO RECOMENDADO", "BAJO"),
]


# ══════════════════════════════════════════════════════════════
# UTILIDADES
# ══════════════════════════════════════════════════════════════

def normalizar(texto: str) -> str:
    """Minúsculas, sin tildes, solo alfanumérico y espacios."""
    texto = texto.lower()
    texto = "".join(
        c for c in unicodedata.normalize("NFD", texto)
        if unicodedata.category(c) != "Mn"
    )
    texto = re.sub(r"[^a-z0-9\s]", " ", texto)
    return re.sub(r"\s+", " ", texto).strip()


def categorizar(puntaje: float) -> tuple[str, str]:
    for umbral, cat, nivel in CATEGORIAS:
        if puntaje >= umbral:
            return cat, nivel
    return "NO RECOMENDADO", "BAJO"


# ══════════════════════════════════════════════════════════════
# EXTRACCIÓN DE TEXTO
# ══════════════════════════════════════════════════════════════

def extraer_texto_pdf(ruta: str) -> str:
    try:
        from pdfminer.high_level import extract_text
        return extract_text(ruta) or ""
    except Exception as e:
        return f"ERROR_PDF: {e}"


def extraer_texto_docx(ruta: str) -> str:
    try:
        from docx import Document
        doc = Document(ruta)
        partes = [p.text for p in doc.paragraphs]
        # Incluir también tablas (experiencia suele estar en tablas)
        for tabla in doc.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    partes.append(celda.text)
        return "\n".join(partes)
    except Exception as e:
        return f"ERROR_DOCX: {e}"


def extraer_texto(ruta: str) -> str:
    ext = os.path.splitext(ruta)[1].lower()
    if ext == ".pdf":
        return extraer_texto_pdf(ruta)
    elif ext in (".docx", ".doc"):
        return extraer_texto_docx(ruta)
    else:
        try:
            with open(ruta, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception:
            return ""


# ══════════════════════════════════════════════════════════════
# MÓDULO 1 — SIMILITUD CON EL PERFIL (30 %)
# ══════════════════════════════════════════════════════════════

def calcular_similitud_perfil(texto_cv: str, perfil: str) -> float:
    """
    TF-IDF coseno entre el CV y el perfil requerido.
    Usa bigramas para capturar frases como 'recursos humanos'.
    """
    if not perfil.strip():
        return 0.0
    try:
        from sklearn.feature_extraction.text import TfidfVectorizer
        from sklearn.metrics.pairwise import cosine_similarity

        vec = TfidfVectorizer(
            analyzer="word",
            ngram_range=(1, 2),
            min_df=1,
            sublinear_tf=True,   # reduce el peso de palabras muy frecuentes
        )
        matriz = vec.fit_transform([normalizar(texto_cv), normalizar(perfil)])
        sim = cosine_similarity(matriz[0], matriz[1])[0][0]
        return round(float(sim) * 100, 2)
    except Exception:
        return 0.0


# ══════════════════════════════════════════════════════════════
# MÓDULO 2 — HABILIDADES TÉCNICAS (25 %)
# ══════════════════════════════════════════════════════════════

def calcular_habilidades(texto_cv: str, keywords_str: str) -> dict:
    """
    Para cada keyword: verifica presencia Y cuenta frecuencia.
    Score = (encontradas / total) * 100, con bonus por frecuencia.
    """
    texto_norm = normalizar(texto_cv)
    palabras = [normalizar(k.strip()) for k in keywords_str.split(",") if k.strip()]

    if not palabras:
        return {
            "pct": 0.0,
            "encontradas": [],
            "total": 0,
            "frecuencias": {},
        }

    encontradas = []
    frecuencias = {}
    for kw in palabras:
        if not kw:
            continue
        # Contar cuántas veces aparece
        freq = len(re.findall(r'\b' + re.escape(kw) + r'\b', texto_norm))
        if freq > 0:
            encontradas.append(kw)
            frecuencias[kw] = freq

    base_pct = (len(encontradas) / len(palabras)) * 100

    # Bonus por alta frecuencia (cap 10 puntos extra)
    if encontradas:
        freq_total = sum(frecuencias.values())
        bonus = min(10, freq_total * 0.5)
    else:
        bonus = 0

    pct = round(min(100, base_pct + bonus), 2)

    return {
        "pct": pct,
        "encontradas": encontradas,
        "total": len(palabras),
        "frecuencias": frecuencias,
    }


# ══════════════════════════════════════════════════════════════
# MÓDULO 3 — EXPERIENCIA LABORAL (20 %)
# ══════════════════════════════════════════════════════════════

def calcular_experiencia(texto_cv: str) -> dict:
    """
    Extrae años de experiencia mencionados explícitamente.
    Patrones: '5 años', '3 años de experiencia', '10+ años', etc.
    También detecta rangos de fechas: 2015-2023 = 8 años.
    """
    texto_norm = normalizar(texto_cv)

    anios_encontrados = []

    # Patrón 1: "N años" mencionados directamente
    patron_directo = re.findall(
        r'(\d{1,2})\s*(?:\+)?\s*(?:años|anos|year[s]?)\s*(?:de\s+(?:experiencia|trabajo|trayectoria))?',
        texto_norm
    )
    for match in patron_directo:
        val = int(match)
        if 1 <= val <= 50:   # filtrar valores absurdos
            anios_encontrados.append(val)

    # Patrón 2: Rangos de fechas laborales (ej: 2018 - 2023, 2018–actualidad)
    patron_rango = re.findall(
        r'(20\d{2}|19\d{2})\s*[-–—]\s*(20\d{2}|19\d{2}|actualidad|presente|actual|current)',
        texto_norm
    )
    anio_actual = 2026
    for inicio, fin in patron_rango:
        try:
            a_inicio = int(inicio)
            a_fin    = anio_actual if fin in ('actualidad', 'presente', 'actual', 'current') else int(fin)
            dur = a_fin - a_inicio
            if 0 < dur <= 45:
                anios_encontrados.append(dur)
        except ValueError:
            pass

    # Tomar el máximo encontrado (representa la experiencia total)
    max_anios = max(anios_encontrados) if anios_encontrados else 0

    # Tabla de puntaje
    if max_anios == 0:
        pct = 0.0
    elif max_anios <= 2:
        pct = 20.0
    elif max_anios <= 5:
        pct = 60.0
    elif max_anios <= 9:
        pct = 80.0
    else:
        pct = 100.0

    return {
        "pct":   pct,
        "anios": max_anios,
        "todos_detectados": sorted(set(anios_encontrados), reverse=True)[:5],
    }


# ══════════════════════════════════════════════════════════════
# MÓDULO 4 — NIVEL EDUCATIVO (15 %)
# ══════════════════════════════════════════════════════════════

def calcular_educacion(texto_cv: str) -> dict:
    """
    Detecta el nivel académico más alto presente en el CV.
    """
    texto_norm = normalizar(texto_cv)
    mejor_nivel = None
    mejor_peso  = 0

    for termino, peso in NIVELES_EDUCATIVOS.items():
        if re.search(r'\b' + re.escape(termino) + r'\b', texto_norm):
            if peso > mejor_peso:
                mejor_peso  = peso
                mejor_nivel = termino

    # El pct es directamente el peso del nivel (ya está en escala 0-100)
    return {
        "pct":   float(mejor_peso),
        "nivel": mejor_nivel or "no detectado",
    }


# ══════════════════════════════════════════════════════════════
# MÓDULO 5 — COMPLETITUD DEL CV (10 %)
# ══════════════════════════════════════════════════════════════

SECCIONES_CV = {
    "experiencia": [
        "experiencia", "trabajo", "empleo", "cargo", "empresa",
        "puesto", "ocupacion", "trayectoria",
    ],
    "educacion": [
        "educacion", "formacion", "universidad", "instituto", "titulo",
        "estudios", "academica", "carrera", "grado",
    ],
    "habilidades": [
        "habilidades", "competencias", "skills", "conocimientos",
        "capacidades", "herramientas", "tecnologias",
    ],
    "datos_personales": [
        "dni", "correo", "email", "telefono", "celular",
        "direccion", "linkedin", "github",
    ],
    "logros": [
        "logros", "logro", "premio", "reconocimiento", "certificado",
        "certificacion", "publicacion", "proyecto",
    ],
    "referencias": [
        "referencia", "referencias", "recomendacion",
    ],
    "idiomas": list(IDIOMAS_CONOCIDOS.keys()),
}

def calcular_completitud(texto_cv: str) -> dict:
    """Detecta qué secciones clave tiene el CV."""
    texto_norm = normalizar(texto_cv)
    presentes  = {}

    for seccion, palabras in SECCIONES_CV.items():
        presentes[seccion] = any(
            re.search(r'\b' + re.escape(p) + r'\b', texto_norm)
            for p in palabras
        )

    total    = len(presentes)
    presentes_count = sum(presentes.values())
    pct      = round((presentes_count / total) * 100, 2) if total else 0.0

    return {
        "pct":      pct,
        "secciones": presentes,
    }


# ══════════════════════════════════════════════════════════════
# MÓDULO EXTRA — IDIOMAS
# ══════════════════════════════════════════════════════════════

def detectar_idiomas(texto_cv: str) -> list[str]:
    texto_norm = normalizar(texto_cv)
    encontrados = []
    for clave, nombre in IDIOMAS_CONOCIDOS.items():
        if re.search(r'\b' + re.escape(clave) + r'\b', texto_norm):
            if nombre not in encontrados:
                encontrados.append(nombre)
    return encontrados


# ══════════════════════════════════════════════════════════════
# GENERADOR DE OBSERVACIONES
# ══════════════════════════════════════════════════════════════

def generar_observaciones(
    puntaje: float,
    categoria: str,
    habilidades: dict,
    experiencia: dict,
    educacion: dict,
    idiomas: list,
) -> str:
    partes = []

    # Veredicto principal
    if categoria == "EXCELENTE":
        partes.append("Perfil altamente compatible con el puesto.")
    elif categoria == "APROBADO":
        partes.append("Perfil compatible con el puesto requerido.")
    elif categoria == "A CONSIDERAR":
        partes.append("Perfil parcialmente compatible. Requiere evaluación adicional.")
    else:
        partes.append("Perfil con baja coincidencia con el puesto.")

    # Experiencia
    anios = experiencia["anios"]
    if anios == 0:
        partes.append("No se detectaron años de experiencia.")
    elif anios <= 2:
        partes.append(f"Experiencia detectada: {anios} año(s) — nivel inicial.")
    elif anios <= 5:
        partes.append(f"Experiencia detectada: {anios} años — nivel medio.")
    else:
        partes.append(f"Experiencia detectada: {anios} años — nivel senior.")

    # Educación
    nivel = educacion["nivel"]
    if nivel != "no detectado":
        partes.append(f"Formación: {nivel.capitalize()}.")
    else:
        partes.append("No se detectó nivel educativo.")

    # Habilidades
    enc = len(habilidades["encontradas"])
    tot = habilidades["total"]
    if tot > 0:
        if enc == 0:
            partes.append("No se encontraron habilidades clave del perfil.")
        elif enc == tot:
            partes.append(f"Todas las habilidades requeridas detectadas ({enc}/{tot}).")
        else:
            partes.append(f"Habilidades detectadas: {enc} de {tot}.")

    # Idiomas
    if idiomas:
        partes.append(f"Idiomas: {', '.join(idiomas)}.")

    return " ".join(partes)


# ══════════════════════════════════════════════════════════════
# LEER ARGUMENTOS
# ══════════════════════════════════════════════════════════════

def leer_argumentos() -> tuple[str, str, str]:
    args = sys.argv[1:]

    if args and args[0] == "--args-file":
        if len(args) < 2:
            print(json.dumps({"error": "Se esperaba la ruta del archivo de argumentos."}))
            sys.exit(1)
        path = args[1]
        if not os.path.exists(path):
            print(json.dumps({"error": f"Archivo de argumentos no encontrado: {path}"}))
            sys.exit(1)
        with open(path, "r", encoding="utf-8-sig") as f:
            datos = json.load(f)
        return (
            datos.get("ruta_cv",  ""),
            datos.get("perfil",   ""),
            datos.get("keywords", ""),
        )
    else:
        if len(args) < 3:
            print(json.dumps({"error": "Uso: analizar_cv.py <ruta_cv> <perfil> <keywords>"}))
            sys.exit(1)
        return args[0], args[1], args[2]


# ══════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════

def main():
    ruta_cv, perfil, keywords = leer_argumentos()

    if not ruta_cv or not os.path.exists(ruta_cv):
        print(json.dumps({"error": f"Archivo no encontrado: {ruta_cv}"}))
        sys.exit(1)

    # ── Extraer texto ─────────────────────────────────────
    texto = extraer_texto(ruta_cv)

    if texto.startswith("ERROR_"):
        print(json.dumps({"error": texto}))
        sys.exit(1)

    if len(texto.strip()) < 30:
        resultado = {
            "puntaje":      0.0,
            "categoria":    "NO RECOMENDADO",
            "nivel":        "BAJO",
            "observaciones":"El archivo no contiene texto suficiente para analizar.",
            "texto_extraido": texto,
            "detalles": {}
        }
        print(json.dumps(resultado, ensure_ascii=False))
        return

    # ── Calcular cada módulo ──────────────────────────────
    m_perfil      = calcular_similitud_perfil(texto, perfil)
    m_habilidades = calcular_habilidades(texto, keywords)
    m_experiencia = calcular_experiencia(texto)
    m_educacion   = calcular_educacion(texto)
    m_completitud = calcular_completitud(texto)
    idiomas       = detectar_idiomas(texto)

    # ── Puntaje ponderado ─────────────────────────────────
    puntaje = (
        m_perfil                * 0.30 +
        m_habilidades["pct"]    * 0.25 +
        m_experiencia["pct"]    * 0.20 +
        m_educacion["pct"]      * 0.15 +
        m_completitud["pct"]    * 0.10
    )
    puntaje = round(min(100.0, max(0.0, puntaje)), 2)

    categoria, nivel = categorizar(puntaje)

    observaciones = generar_observaciones(
        puntaje, categoria,
        m_habilidades, m_experiencia, m_educacion, idiomas
    )

    # ── Resultado final ───────────────────────────────────
    # texto_extraido se pasa como campo separado pero limpio (ASCII-safe)
    # para que PHP pueda guardarlo en BD sin problemas de encoding
    texto_limpio = texto.encode('utf-8', errors='replace').decode('utf-8')

    resultado = {
        "puntaje":       puntaje,
        "porcentaje_coincidencia": m_perfil,  # Agregar explícitamente para la BD
        "categoria":     categoria,
        "nivel":         nivel,
        "observaciones": observaciones,
        "texto_extraido": texto_limpio,   # para BD, no para la UI
        "detalles": {
            # Puntajes por módulo (0-100)
            "perfil_pct":          m_perfil,
            "habilidades_pct":     m_habilidades["pct"],
            "experiencia_pct":     m_experiencia["pct"],
            "educacion_pct":       m_educacion["pct"],
            "completitud_pct":     m_completitud["pct"],

            # Datos cualitativos
            "anios_experiencia":        m_experiencia["anios"],
            "nivel_educativo":          m_educacion["nivel"],
            "habilidades_encontradas":  m_habilidades["encontradas"],
            "habilidades_total":        m_habilidades["total"],
            "habilidades_frecuencias":  m_habilidades["frecuencias"],
            "idiomas":                  idiomas,
            "secciones":                m_completitud["secciones"],
        }
    }

    # Imprimir SOLO el JSON en stdout (pdfminer va a stderr)
    print(json.dumps(resultado, ensure_ascii=True))  # ASCII-safe para PHP


if __name__ == "__main__":
    main()
